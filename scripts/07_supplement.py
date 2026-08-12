#!/usr/bin/env python3
"""Build supplement.docx for the C71 age study."""
import csv, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"
HDR = "D9E2F3"
RT = "Age and hospital care in malignant glioma"
BANDS = ["18-29", "30-39", "40-49", "50-54", "55-59", "60-64", "65-74", "75-79", ">=80"]
DISP = ["18-29", "30-39", "40-49", "50-54", "55-59", "60-64", "65-74", "75-79", "≥80"]

stats = json.load(open("derived_data/statistics.json"))
_meta = json.load(open("study_meta.json"))
STUDY_YEAR = _meta["study_year"]
SUPP = json.load(open("derived_data/suppressed_cells.json"))
union = stats["meta"]["union_by_band"]
UN = [union[b] for b in BANDS]

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT; st.font.size = Pt(11); st.font.color.rgb = BLACK
st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def style_section(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11.69), Inches(8.27)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Inches(8.27), Inches(11.69)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Inches(0.8))
    # footer page number
    fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    run = fp.add_run()
    for t, v in [("begin", None), ("instr", "PAGE"), ("end", None)]:
        if t == "instr":
            e = OxmlElement("w:instrText"); e.set(qn("xml:space"), "preserve"); e.text = v
        else:
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), t)
        run._r.append(e)
    run.font.name = FONT; run.font.size = Pt(11); run.font.color.rgb = BLACK
    section.footer.is_linked_to_previous = False
    # header running title
    hp = section.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.text = ""
    r = hp.add_run(RT); r.italic = True
    r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = BLACK
    section.header.is_linked_to_previous = False


style_section(doc.sections[0], landscape=False)


def para(text, bold=False, italic=False, size=11, after=6, before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format; pf.space_after = Pt(after); pf.space_before = Pt(before)
    pf.line_spacing = 1.0
    if text:
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = BLACK
    return p


def shade(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), HDR); tcpr.append(sh)


def setc(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""; p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = bold
    r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = BLACK


def table(headers, data, widths, total_in, first_left=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.autofit = False
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(int(total_in*1440)))
    tblW.set(qn("w:type"), "dxa"); t._tbl.tblPr.append(tblW)
    for i, c in enumerate(t.rows[0].cells):
        setc(c, headers[i], bold=True, size=8,
             align=WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_left) else WD_ALIGN_PARAGRAPH.CENTER)
        shade(c)
    for drow in data:
        cells = t.add_row().cells
        for i, val in enumerate(drow):
            setc(cells[i], val, size=8,
                 align=WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_left) else WD_ALIGN_PARAGRAPH.CENTER)
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Inches(widths[i])
    return t


def fmt_p(p):
    if p is None:
        return "NA"
    p = float(p)
    return "<0.001" if p < 0.001 else f"{p:.3f}"


# ===========================================================================
# Title
# ===========================================================================
para("Supplementary Material", bold=True, size=14, after=8)
para("Age-stratified hospital care, comorbidity, and discharge outcomes in adults "
     "with malignant glioma in Germany: a nationwide analysis of administrative data",
     italic=True, size=11, after=10)
para("Contents: S1 Codebook and definitions; S2 Methods appendix; "
     "S3 Supplementary tables (S1-S6); S4 Reporting checklist (RECORD); "
     "Supplementary Figure S1 (cohort derivation).", size=11, after=10)

# ===========================================================================
# S1 Codebook
# ===========================================================================
para("S1. Codebook and definitions", bold=True, size=12, after=6)
para("Data source. Aggregate hospital discharge data from the German §21 Hospital "
     "Remuneration Act, published through the InEK DatenBrowser. The unit of "
     "observation is the hospital discharge episode.", size=11, after=6)
para("Diagnosis. Malignant neoplasm of brain, ICD-10-GM code C71 (all subcodes "
     "C71.0-C71.9). The code is topographic and does not encode histology, "
     "molecular status, or WHO grade.", size=11, after=6)
para("Cohorts. (i) Principal-diagnosis cohort: episodes with C71 as the principal "
     "(main) diagnosis (n = 18,621), used for age, sex, in-hospital mortality, and "
     "discharge disposition. (ii) C71-coded population: episodes with C71 as "
     "principal or secondary diagnosis (n = 23,594), the denominator to which the "
     "comorbidity, secondary-diagnosis, and procedure tabulations in the source "
     "export are keyed; used for those analyses only.", size=11, after=6)
para("Age bands. 18-29, 30-39, 40-49, 50-54, 55-59, 60-64, 65-74, 75-79, and 80 "
     "years or older; bands are of unequal width, as provided by the source export. "
     "Trend tests used age-band midpoints as scores (23.5, 34.5, 44.5, 52, 57, 62, "
     "69.5, 77, 85).", size=11, after=6)
para("Discharge disposition. Death in hospital; discharge home; transfer to a "
     "nursing home; transfer to hospice care; transfer to another hospital; other "
     "or unspecified. For the principal-diagnosis cohort, main-diagnosis counts are "
     "available for the first five categories; the residual to the age-band total "
     "was assigned to other or unspecified.", size=11, after=6)
para("Comorbidity. Grouped into Charlson comorbidity index categories from recorded "
     "ICD-10-GM secondary diagnoses. Ten categories with complete tabulation are "
     "reported: other solid tumour (non-C71), malignant haematological disease, "
     "diabetes mellitus, dementia, heart failure, cerebrovascular disease, chronic "
     "pulmonary disease, peripheral arterial occlusive disease, collagenoses, and "
     "renal failure or chronic kidney disease.", size=11, after=6)
para("Procedures (OPS codes). Single codes analysed as exact per-episode "
     "frequencies: 3-200 native cranial CT; 3-800 native cranial MRI; 5-984 "
     "microsurgical technique; 5-989 fluorescence-guided surgery; 5-022.00 external "
     "ventricular drain. Grouped code families analysed as approximate frequencies: "
     "5-988.x intraoperative navigation; 1-511.00/1-511.01 supratentorial "
     "stereotactic biopsy; 8-522.x inpatient high-voltage radiotherapy. Procedure "
     "categories are not mutually exclusive; a single episode may carry several "
     "codes, so percentages are not additive.", size=11, after=6)
para("Privacy suppression. The InEK DatenBrowser aggregates only cells with five "
     "or more cases; cells with fewer than five cases are suppressed (blank in the "
     "export). Suppressed cells are shown as <5 (counts) or <0.X (percentages, an "
     "upper bound) and were treated as zero only in the comorbidity-point sum, "
     "which slightly understates the youngest bands.", size=11, after=10)

# ===========================================================================
# S2 Methods appendix
# ===========================================================================
para("S2. Methods appendix", bold=True, size=12, after=6)
para("Source workbook structure. The export comprises six worksheets: an overview "
     "of case counts, sex, and discharge disposition by age band and diagnosis "
     "position; a sex-by-disposition cross-tabulation; a Charlson comorbidity "
     "tabulation; a secondary-diagnosis frequency table; and two procedure-frequency "
     "tables. All counts were read programmatically; percentages were recomputed "
     "from raw counts.", size=11, after=6)
para("Statistical methods. Proportions are reported with 95% Wilson confidence "
     "intervals. Association between age band and each binary outcome was tested "
     "with the Pearson chi-square test of independence. Monotone trend across the "
     "ordered age bands was tested with the Cochran-Armitage trend test using "
     "age-band midpoint scores; the test statistic is asymptotically standard "
     "normal, and two-sided p-values are reported. Within each family of trend "
     "tests (discharge disposition, comorbidity, and procedures), p-values were "
     "adjusted across outcomes using the Benjamini-Hochberg false discovery rate; "
     "for the five discharge-disposition outcomes the adjusted p-values remained "
     "below 0.001. Effect "
     "sizes are reported as absolute risk differences (oldest versus youngest band) "
     "with Newcombe 95% confidence intervals, odds ratios per decade of age from "
     "grouped-binomial logistic trend models, and Cramer's V for the global "
     "contingency tests (Table S5). Trend results were checked for robustness to "
     "age-band scoring using integer rank scores and alternative open-band scores "
     "(Table S4). Analyses used Python 3.10 with "
     "NumPy; chi-square survival probabilities were computed from the regularised "
     "incomplete gamma function. Because the data are aggregate, no individual-level "
     "adjustment or multivariable modelling was possible.", size=11, after=6)
para("Charlson-weighted comorbidity point score. The mean Charlson-weighted "
     "comorbidity point score per episode within each age band was computed as the "
     "sum over comorbidity categories of the Charlson weight multiplied by the "
     "category frequency, divided by the number of episodes. This is a weighted "
     "sum of comorbidity indicators using Charlson weights, not a patient-level "
     "Charlson index applying the standard comorbidity hierarchy. Because a mean "
     "is a linear "
     "function of the marginal frequencies, it is exact even though co-occurrence "
     "within episodes is unknown; the per-episode score distribution is not "
     "recoverable. Diabetes was split into uncomplicated (weight 1) and complicated "
     "(weight 2) and the second-solid-tumour category into non-metastatic "
     "(weight 2) and metastatic (weight 6) at ICD-code level. Secondary neoplasms "
     "of the brain and other nervous system (C79.3, C79.4) were excluded from the "
     "tumour comorbidity, as in a primary brain-tumour cohort they most plausibly "
     "reflect involvement of the index tumour rather than a separate metastatic "
     "cancer. Age points were omitted. The trend was tested with Poisson regression "
     "using episode counts as exposure. Cells suppressed for privacy were treated "
     "as zero in this point sum, which slightly understates the youngest bands. "
     "The weighted comorbidity points are "
     "strongly overdispersed relative to a Poisson model (Pearson dispersion "
     "statistic 41.5 on 7 degrees of freedom), so the rate ratio is reported with "
     "a quasi-Poisson 95% confidence interval (1.46 per decade; 95% CI 1.33 to "
     "1.59), wider than the naive Poisson interval (1.44 to 1.48).", size=11, after=6)
para("Reproducibility. Analysis code, data-extraction instructions, and an "
     "in-depth study summary are openly available in a public GitHub repository "
     "(https://github.com/rnmcln/inek-c71-age-2026). To respect the terms of the "
     "data source, the repository does not redistribute the InEK data or derived "
     "data tables; all reported numbers can be regenerated from a DatenBrowser "
     "export by running the pipeline. On acceptance, the repository will be "
     "archived on Zenodo and assigned a citable DOI.", size=11, after=6)

para("Supplementary Figure S1", bold=True, size=11, after=2)
para("Cohort derivation and the two denominators used in the analysis.", size=10, after=4)
_figp = doc.add_paragraph(); _figp.alignment = WD_ALIGN_PARAGRAPH.CENTER
_figp.add_run().add_picture("figures/FigS1_cohort.png", width=Inches(5.5))

# ---- landscape section for wide tables ----
doc.add_section(WD_SECTION.NEW_PAGE)
style_section(doc.sections[-1], landscape=True)

para("S3. Supplementary tables", bold=True, size=12, after=6)

# Table S1 comorbidity
para("Table S1", bold=True, size=11, after=2)
para("Coded prevalence of Charlson comorbidity categories by age band in the "
     "C71-coded population (principal or secondary diagnosis; n = 23,594). Values "
     "are percentages of episodes within each age band; the final column gives the "
     "false-discovery-rate-adjusted Cochran-Armitage trend p-value.", size=10, after=6)

com_rows = list(csv.DictReader(open("derived_data/comorbidity_by_age.csv")))
headers = ["Comorbidity"] + DISP + ["Trend p (FDR)"]
data = []
for r in com_rows:
    name = r["comorbidity"]
    counts = [int(round(float(r[b]))) for b in BANDS]
    pcts = []
    for c, u, b in zip(counts, UN, BANDS):
        if SUPP["comorbidity"].get(name, {}).get(b, False):
            pcts.append(f"<{100*5/u:.1f}")
        else:
            pcts.append(f"{100*c/u:.1f}")
    data.append([name] + pcts + [fmt_p(r["p_trend_fdr"])])
cci = json.load(open("derived_data/charlson_index.json"))
cci_vals = [f"{cci['mean_cci'][b]:.2f}" for b in BANDS]
data.append(["Charlson-weighted comorbidity point score (a)"] + cci_vals + [fmt_p(cci["p_trend"])])
w = [2.6] + [0.72]*9 + [0.95]
table(headers, data, w, sum(w))
para("(a) Final row: mean Charlson-weighted comorbidity point score per episode (a "
     "weighted score, not a percentage), computed from the marginal category "
     "frequencies with ICD-code-level weighting for complicated diabetes and "
     "metastatic disease and no age points. Secondary neoplasms of the brain and "
     "other nervous system (C79.3, C79.4) were not counted as a metastatic-tumour "
     "comorbidity, as in a primary brain-tumour cohort they most plausibly reflect "
     "CNS involvement of the index tumour. The trend in the comorbidity-point rate "
     "was tested by Poisson regression (rate ratio 1.46 per decade; quasi-Poisson "
     "95% CI 1.33 to 1.59, accounting for overdispersion). A "
     "patient-level score distribution cannot be derived from aggregate data.",
     size=9, after=6, before=2)

# Table S2 procedures
para("Table S2", bold=True, size=11, after=2)
para("Coding frequency of selected diagnostic and therapeutic procedures by age "
     "band in the C71-coded population (n = 23,594). Values are percentages of "
     "episodes within each age band. Grouped categories (navigation, stereotactic "
     "biopsy, radiotherapy) are approximate because a single episode may carry "
     "several codes. The final column gives the FDR-adjusted trend p-value.",
     size=10, after=6)
proc_rows = list(csv.DictReader(open("derived_data/procedures_by_age.csv")))
headers2 = ["Procedure (OPS)"] + DISP + ["Trend p (FDR)"]
data2 = []
for r in proc_rows:
    counts = [int(round(float(r[b]))) for b in BANDS]
    pcts = [f"{100*c/u:.1f}" for c, u in zip(counts, UN)]
    data2.append([r["procedure"]] + pcts + [fmt_p(r["p_trend_fdr"])])
w2 = [2.9] + [0.66]*9 + [0.9]
table(headers2, data2, w2, sum(w2))
para("", after=4)

# Table S3 union denominators
para("Table S3", bold=True, size=11, after=2)
para(f"Number of discharge episodes by age band in each cohort ({STUDY_YEAR} data year).", size=10, after=6)
h3 = ["Cohort"] + DISP + ["Total"]
prin = stats["meta"]["principal_by_band"]
row_p = ["C71 principal diagnosis"] + [str(prin[b]) for b in BANDS] + [str(sum(prin.values()))]
row_u = ["C71 principal or secondary"] + [str(union[b]) for b in BANDS] + [str(sum(union.values()))]
w3 = [2.8] + [0.72]*9 + [0.9]
table(h3, [row_p, row_u], w3, sum(w3))

# ---- Table S4: Cochran-Armitage score sensitivity ----
import json as _json
_eff = _json.load(open("derived_data/effects_sensitivity.json"))
doc.add_page_break()
para("Table S4", bold=True, size=11, after=2)
para("Sensitivity of Cochran-Armitage trend p-values to the scoring of age bands. "
     "Columns give the two-sided, unadjusted (raw) trend p-value under the age-band "
     "midpoint scores "
     "used in the primary analysis, under integer rank scores (1-9), and under two "
     "alternative scores for the open-ended oldest band (82.5 and 90 years). The "
     "direction and significance of every trend are preserved across scorings; the "
     "only borderline outcome is fluorescence-guided surgery, whose weak trend is "
     "not robust and is not interpreted in the main text.", size=10, after=6)
_sh = {"Comorbidity: ": "", "Procedure: ": ""}
def _short(k):
    for a, b in _sh.items():
        if k.startswith(a):
            return k[len(a):]
    return k
hS4 = ["Outcome", "Midpoint (primary)", "Rank scores", "Oldest = 82.5", "Oldest = 90"]
dS4 = []
for k, v in _eff["ca_sensitivity"].items():
    dS4.append([_short(k), fmt_p(v["midpoint"]["p"]), fmt_p(v["rank"]["p"]),
                fmt_p(v["oldest_82.5"]["p"]), fmt_p(v["oldest_90"]["p"])])
wS4 = [4.0, 1.3, 1.1, 1.1, 1.0]
table(hS4, dS4, wS4, sum(wS4))
para("", after=6)

# ---- Table S5: effect sizes ----
para("Table S5", bold=True, size=11, after=2)
para("Effect sizes for the principal-diagnosis cohort outcomes. Risk difference is "
     "the absolute difference in the outcome proportion between the oldest (>=80) "
     "and youngest (18-29) bands, in percentage points, with a Newcombe 95% "
     "confidence interval. The odds ratio per decade of age is from a "
     "grouped-binomial logistic trend model. For the global age-by-category "
     "associations, Cramer's V was 0.06 (sex) and 0.11 (discharge disposition), "
     "indicating that the associations, although highly significant, are modest in "
     "standardised magnitude.", size=10, after=6)
hS5 = ["Outcome", "Youngest %", "Oldest %", "Risk difference, pp (95% CI)",
       "OR per decade (95% CI)"]
dS5 = []
for k, v in _eff["effect_sizes"].items():
    rd = v["risk_diff"] * 100; lo = v["rd_ci"][0] * 100; hi = v["rd_ci"][1] * 100
    orr = v["or_per_decade"]; olo = v["or_ci"][0]; ohi = v["or_ci"][1]
    dS5.append([k, f"{100*v['p_youngest']:.1f}", f"{100*v['p_oldest']:.1f}",
                f"{rd:+.1f} ({lo:+.1f} to {hi:+.1f})", f"{orr:.2f} ({olo:.2f}-{ohi:.2f})"])
wS5 = [2.9, 1.0, 1.0, 2.6, 2.0]
table(hS5, dS5, wS5, sum(wS5))

# ---- Table S6: absolute older-age service burden ----
doc.add_page_break()
para("Table S6", bold=True, size=11, after=2)
para("Absolute service burden in older age groups (principal-diagnosis cohort). "
     "For episodes in patients aged 65 or older and 75 or older, the number and "
     "within-group percentage of episodes ending in each outcome.", size=10, after=6)
_ob = _eff["older_age_burden"]
hS6 = ["Age group", "Episodes, n (% of cohort)", "In-hospital death, n (%)",
       "Nursing-home transfer, n (%)", "Hospice transfer, n (%)",
       "Inter-hospital transfer, n (%)"]
dS6 = []
for thr, lab in [(">=65", "\u2265 65 years"), (">=75", "\u2265 75 years")]:
    rr = _ob[thr]
    def _c(o):
        return f"{rr[o]['n']} ({rr[o]['pct_of_group']:.1f})"
    dS6.append([lab, f"{rr['n_episodes']} ({rr['pct_of_cohort']:.1f})",
                _c('death'), _c('nursing_home'), _c('hospice'), _c('transfer_hospital')])
wS6 = [1.2, 1.9, 1.5, 1.7, 1.5, 1.7]
table(hS6, dS6, wS6, sum(wS6))

# ---- S4. Reporting checklist (RECORD) ----
doc.add_section(WD_SECTION.NEW_PAGE)
style_section(doc.sections[-1], landscape=False)
para("S4. Reporting checklist (RECORD)", bold=True, size=12, after=6)
para("Items from the REporting of studies Conducted using Observational "
     "Routinely-collected health Data (RECORD) statement that extend STROBE, and "
     "where each is addressed. Core STROBE items are addressed in the main text.",
     size=10, after=6)
_record = [
  ("RECORD 1.1", "Study design and data source indicated in title/abstract", "Title; Abstract (Methods)"),
  ("RECORD 6.1", "Population selection: codes, algorithms, diagnosis position", "Methods (Study population); Suppl. S1"),
  ("RECORD 6.2", "Validation of codes used to select the population", "Not chart-validated (Limitations)"),
  ("RECORD 6.3", "Data linkage across databases or records", "Not applicable (aggregate export)"),
  ("RECORD 7.1", "Complete list of codes and algorithms used", "Suppl. S1 (ICD-10-GM, OPS codebook)"),
  ("RECORD 12.1", "Data-cleaning methods", "Methods (Statistical analysis); Suppl. S2"),
  ("RECORD 12.2", "Linkage methods and quality", "Not applicable"),
  ("RECORD 12.3", "Derivation of the study population size", "Methods; Table S3; Denominators note"),
  ("RECORD 13.1", "Selection of included episodes (flow)", "Table S3; Fig. S1 (cohort flow)"),
  ("RECORD 19.1", "Implications of using non-research data", "Limitations"),
  ("RECORD 22.1", "Access to protocol, raw data, and code", "Data availability; GitHub repository"),
]
table(["RECORD item", "Recommendation", "Location in this report"],
      [list(x) for x in _record], [1.1, 3.3, 2.2], 6.6)

doc.save("supplement.docx")
print("saved supplement.docx")
