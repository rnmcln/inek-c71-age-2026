#!/usr/bin/env python3
"""Build tables.docx (Table 1, Table 2) for the C71 age study."""
import csv, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"
HDR_SHADE = "D9E2F3"
RUNNING_TITLE = "Age and hospital care in malignant glioma"

rows = {r["age_band"]: r for r in csv.DictReader(
    open("derived_data/principal_cohort_by_age.csv"))}
BANDS = ["18-29", "30-39", "40-49", "50-54", "55-59", "60-64", "65-74", "75-79", ">=80"]
DISP = {"18-29": "18-29", "30-39": "30-39", "40-49": "40-49", "50-54": "50-54",
        "55-59": "55-59", "60-64": "60-64", "65-74": "65-74", "75-79": "75-79",
        ">=80": "≥80"}
N = sum(int(rows[b]["n_principal"]) for b in BANDS)
_meta = json.load(open("study_meta.json"))
STUDY_YEAR = _meta["study_year"]
SUPP = json.load(open("derived_data/suppressed_cells.json"))
WIDTH = {"18-29": 12, "30-39": 10, "40-49": 10, "50-54": 5, "55-59": 5,
         "60-64": 5, "65-74": 10, "75-79": 5, ">=80": None}

doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT; style.font.size = Pt(11); style.font.color.rgb = BLACK
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(0.8))


def add_field_footer(section):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for t, v in [("begin", None), ("instr", "PAGE"), ("end", None)]:
        if t == "instr":
            e = OxmlElement("w:instrText"); e.set(qn("xml:space"), "preserve"); e.text = v
        else:
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), t)
        run._r.append(e)
    run.font.name = FONT; run.font.size = Pt(11); run.font.color.rgb = BLACK


def add_header(section):
    p = section.header.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(RUNNING_TITLE); r.italic = True
    r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = BLACK


add_field_footer(sec); add_header(sec)


def para(text, bold=False, italic=False, size=11, after=6, before=0):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(after); pf.space_before = Pt(before); pf.line_spacing = 1.0
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = BLACK
    return p


def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcpr.append(sh)


def set_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = bold
    r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = BLACK


def make_table(headers, data_rows, widths, total_width_in):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    # column widths
    tblPr = t._tbl.tblPr
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(int(total_width_in*1440)))
    tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    for i, c in enumerate(t.rows[0].cells):
        set_cell(c, headers[i], bold=True, size=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT)
        shade(c, HDR_SHADE)
    for drow in data_rows:
        cells = t.add_row().cells
        for i, val in enumerate(drow):
            set_cell(cells[i], val, bold=(drow is data_rows[-1]), size=9,
                     align=WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT)
    # apply widths to every cell
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Inches(widths[i])
    return t


def pc(k, n):
    return f"{k} ({100*k/n:.1f})"


# ===========================================================================
# Table 1 — cohort characteristics by age
# ===========================================================================
para("Table 1", bold=True, size=11, after=2)
para(f"Age and sex structure of the principal-diagnosis cohort (malignant glioma "
     f"as principal hospital diagnosis; {STUDY_YEAR} data year; n = 18,621).",
     italic=False, size=11, after=6)

h1 = ["Age band (years)", "Discharge episodes, n (%)", "Per year of band width",
      "Male, n (%)", "Female, n (%)"]
d1 = []
for b in BANDS:
    n = int(rows[b]["n_principal"]); m = int(rows[b]["male"]); f = int(rows[b]["female"])
    pyw = f"{n/WIDTH[b]:.0f}" if WIDTH[b] else "n/a (open)"
    d1.append([DISP[b], f"{n} ({100*n/N:.1f})", pyw, pc(m, n), pc(f, n)])
tot_m = sum(int(rows[b]["male"]) for b in BANDS)
tot_f = sum(int(rows[b]["female"]) for b in BANDS)
d1.append(["All ages", f"{N} (100.0)", "-", pc(tot_m, N), pc(tot_f, N)])
make_table(h1, d1, [1.5, 2.1, 1.8, 1.9, 1.9], 9.2)
para("Percentages for episodes are column percentages of the total cohort; sex "
     "percentages are row percentages within each age band. \"Per year of band "
     "width\" divides the episode count by the band width in years (the open-ended "
     "oldest band has no defined width); because bands are unequal, counts index "
     "hospital workload rather than age-specific risk or incidence. The distribution "
     "of sex differed across age bands (chi-square = 67.3, 8 df, p < 0.001; "
     "Cramer's V = 0.06), with a nonlinear male fraction that was lowest in the "
     "oldest band (Cochran-Armitage trend p < 0.001).", size=9, after=4, before=6)

doc.add_page_break()

# ===========================================================================
# Table 2 — discharge disposition by age
# ===========================================================================
para("Table 2", bold=True, size=11, after=2)
para("Discharge disposition by age band in the principal-diagnosis cohort, n (%). "
     "Percentages are within-age-band proportions of all episodes.", size=11, after=6)

h2 = ["Age band (years)", "Death in hospital", "Discharged home", "Nursing home",
      "Hospice care", "Other hospital", "Other/unspecified"]
keymap = [("death",), ("home",), ("nursing_home",), ("hospice",),
          ("transfer_hospital",), ("other_residual",)]
d2 = []
for b in BANDS:
    n = int(rows[b]["n_principal"])
    cells = [DISP[b]]
    for (k,) in keymap:
        if k != "other_residual" and SUPP["disposition"].get(k, {}).get(b, False):
            cells.append("<5")
        else:
            v = int(rows[b][k])
            cells.append(f"{v} ({100*v/n:.1f})")
    d2.append(cells)
overall = ["All ages"]
for (k,) in keymap:
    v = sum(int(rows[b][k]) for b in BANDS)
    overall.append(f"{v} ({100*v/N:.1f})")
d2.append(overall)
make_table(h2, d2, [1.5, 1.4, 1.4, 1.2, 1.2, 1.3, 1.4], 9.4)
para("All disposition categories showed a significant gradient with age "
     "(Pearson chi-square and Cochran-Armitage trend, all p < 0.001). The overall "
     "age-by-disposition association was modest in standardised terms (Cramer's "
     "V = 0.11). "
     "In-hospital mortality and transfers to nursing home, hospice, and other "
     "hospitals increased with age, while discharge home decreased. Cells with "
     "fewer than five episodes are suppressed in the source data and shown as <5; "
     "the other/unspecified category is derived as the residual to the age-band "
     "total.", size=9, after=4, before=6)

doc.save("tables.docx")
print("saved tables.docx; N =", N)
